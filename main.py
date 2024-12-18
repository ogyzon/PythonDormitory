'''Жоров Евгений Александрович
гр. 10701123
@vanasokolov844@gmail.com
Курсовой проект по дисциплине "Языки программирования"
Минск 2024'''


'''копия в ef'''

'''Учет дежурств в общежитии'''

'''Окно сортировки немного лишнее. Лучше сделать мини менюшку с сортировкой. Реализовать поиск(к примеру поиск по датам будет уместно)'''

import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog
from openpyxl import Workbook
from openpyxl import load_workbook
from docx import Document
from openpyxl.styles import Font
import threading
import time
import os
import telebot
import requests
import re

# Класс для Label
class MyLabel(tk.Label):

    def __init__(self, master, text=None, x=0, y=0, font_size=12, font_weight="normal", justify="left", bg="#f0f0f0",
                 image=None, **kwargs):
        font = ("Montserrat", font_size, font_weight)

        super().__init__(master, text=text, font=font, justify=justify, bg=bg, image=image, **kwargs)
        self.place(x=x, y=y)


# Класс для кнопки
class MyButton(tk.Button):

    def __init__(self, master, text, width, height, fg="#000000", bg="#f0f0f0", image=None, compound=None, font_size=9,
                 font_weight="normal", command=None, x=0, y=0, **kwargs):
        font = ("Montserrat", font_size, font_weight)

        super().__init__(master, text=text, font=font, width=width, height=height, fg=fg, bg=bg, image=image,
                         compound=compound, command=command, **kwargs)
        self.place(x=x, y=y)

class User():

    def __init__(self, username, role):

        self.username = username
        self.role = role

#Класс для отслеживания неактивности в течение 1 мин
class Inactivity:
    def __init__(self, frame):
        self.frame = frame
        self.lastActivityTime = time.time()
        self.is_active = False  # Флаг для управления активностью
        self.setupBindings()  # Обработчик событий

    def resetTime(self, event):
        self.lastActivityTime = time.time()

    def setupBindings(self):
        self.frame.bind_all('<Any-KeyPress>', self.resetTime)
        self.frame.bind_all('<Any-Button>', self.resetTime)

    def startCheck(self):
        self.is_active = True
        self.checkInactivity()

    def stopCheck(self):
        self.is_active = False

    def checkInactivity(self):
        if not self.is_active:
            return

        currentTime = time.time()
        timeDuration = currentTime - self.lastActivityTime

        if timeDuration >= 60:
            self.showWarning()


        # Проверка на неактивность каждые 1000 миллисекунд (1 секунда)
        self.frame.after(1000, self.checkInactivity)

    def showWarning(self):
        messagebox.showwarning("Бездействие", "Вы бездействуете 60 секунд")


# Класс приложения
class MyApp(tk.Tk):

    def __init__(self):

        super().__init__()

        self.title("Учет дежурств")
        self.geometry('800x850')
        self.geometry('+500+100')
        self.resizable(width=False, height=False)

        self.user = None
        self.roles = {"commandant" : "комендант",
                      "student" : "студент"}

        self.protocol("WM_DELETE_WINDOW", self.onClosing)  # Обработчик закрытия окна
        self.telegramBot = MyTelegramBot(self)
        self.chatID = 990537084

        self.filePath = "excelFiles//Book1.xlsx"

        if not os.path.exists(self.filePath):
            self.excelObjectFirst = WorkExcel(self, self.filePath)

        self.selected_index = None

        # Словарь для хранения записанных дат
        self.dictZapisanye = {}

        # Для перевода из слова в номер месяца
        self.DictForMonthes = {"Январь": '01', "Февраль": "02", "Март": "03", "Апрель": "04",
                               "Май": "05", "Июнь": "06", "Июль": "07", "Август": "08",
                               "Сентябрь": "09", "Октябрь": "10", "Ноябрь": "11", "Декабрь": "12"}

        self.stringDay = ""
        self.stringMonthNum = ""
        self.FullConcatenate = ""

        self.concatenateString = ""
        self.monthUser = ""

        self.fullStringSort = ""

        # Для сортировки и Excel
        self.recordBlocks = []
        self.recordRooms = []
        self.recordSecondNames = []
        self.recordFirstNames = []
        self.recordDays = []
        self.recordMonthes = []
        self.recordTelNumber = []
        self.recordTimes = []
        self.recordFullStrings = []
        self.recordDates = []
        self.recordYears = []
        self.studentRecords = []

        self.dropDownVisibleSort = False
        self.dropDownVisibleSearch = False

        # Фрейм самого первого Splash окна
        self.frame1 = tk.Frame(self, width=800, height=800)
        self.frame1.place(x=0, y=0)

        self.inactivity = Inactivity(self.frame1)
        self.inactivity.startCheck()

        self.sortWindowInstance = None

        #Фрейм окна аутентификации
        self.frameAuthentification = tk.Frame(self, width=800, height = 800)

        # Фрейм второго (основного окна)
        self.frameMain = tk.Frame(self, width=800, height=800)

        # Фрейм для окна об авторе
        self.frameAuthor = tk.Frame(self, width=800, height=850)

        self.frameAbtProgramm = tk.Frame(self, width=800, height=850)

        self.canvasAbtProgramm = tk.Canvas(self.frameAbtProgramm, width=300, height=850, bg='#3469ED')
        self.canvasAbtProgramm.place(x=0, y=0)

        self.canvasAbtAuthor = tk.Canvas(self.frameAuthor, width=300, height=550, bg='#e3e2e1')
        self.canvasAbtAuthor.place(x=40, y=150)

        self.canvasAbtAuthor2 = tk.Canvas(self.frameAuthor, width=230, height=230)
        self.canvasAbtAuthor2.place(x=75, y=420)

        self.canvasAbtAuthor3 = tk.Canvas(self.frameAuthor, width=170, height=3, bg='#3469ed')
        self.canvasAbtAuthor3.place(x=520, y=163)

        self.canvas1 = tk.Canvas(self.frameMain, width=760, height=670, bg='#f1f1f1', borderwidth=2, relief='solid')
        self.canvas1.place(x=0, y=50)
        self.canvas1.bind('<Button-1>', self.closeFrame)

        self.canvas2 = tk.Canvas(self.frameMain, width=158, height=21, bg='#1F1633')
        self.canvas2.place(x=305, y=20)

        self.canvas3 = tk.Canvas(self.frameMain, width=158, height=21, bg='#3469ed')
        self.canvas3.place(x=293, y=8)

        self.canvas4 = tk.Canvas(self.canvas1,width=612, height=393, bg='#3469ed')
        self.canvas4.place(x=18, y=219)

        self.canvas5 = tk.Canvas(self.canvas1, width=605, height=391, bg='#1F1633')
        self.canvas5.place(x=18, y=213)

        self.canvas6 = tk.Canvas(self.frameMain, width=120, height=3, bg='#3469ed')
        self.canvas6.place(x=355, y=775)

        self.canvas7 = tk.Canvas(self.frameMain, width=2, height=15, bg='#545454')
        self.canvas7.place(x=120, y=61)

        self.frameText = tk.Frame(self.frameMain)
        self.frameText.place(x=20, y=250, width=600, height=400)

        self.textPlace = tk.Text(self.frameText, state='normal', font=("Montserrat", 8, 'bold'), borderwidth=2)
        self.textPlace.pack(side="left", fill="both", expand=True)
        self.textPlace.insert(tk.END, f"Фамилия\t\t  Имя\t\tБлок\t\tTелефон\t\t   Дата\t\tВремя\n\n")

        self.scrollBar = tk.Scrollbar(self.frameText, command=self.textPlace.yview)
        self.scrollBar.pack(side="right", fill="y")
        self.textPlace.config(yscrollcommand=self.scrollBar.set)
        self.textPlace.config(state='disabled')
        self.textPlace.bind('<ButtonRelease-1>', self.onTextClick)

        # Лейблы с текстом Splash окна
        self.label1 = MyLabel(self.frame1, text='Белорусский национальный технический университет', x=195, y=20)

        self.label2 = MyLabel(self.frame1, text='Факультет информационных технологий и робототехники', x=175, y=50)

        self.label3 = MyLabel(self.frame1, text='Кафедра программного обеспечения информационных систем и технологий',
                              x=120, y=80)

        self.label4 = MyLabel(self.frame1, text='Курсовой проект', font_size=16, font_weight='bold', x=300, y=160)

        self.label5 = MyLabel(self.frame1, text='По дисциплине "Языки программирования"', font_size=16,
                              font_weight='bold', x=160, y=200)

        self.label6 = MyLabel(self.frame1, text='Выполнил: студент группы 10701123', font_size=12, font_weight='bold',
                              x=340, y=360)

        self.label7 = MyLabel(self.frame1, text='Жоров Евгений Александрович', font_size=12, font_weight='bold', x=340,
                              y=390)

        self.label8 = MyLabel(self.frame1, text='Преподаватель: к.ф.-м.н., доц.', font_size=12, font_weight='bold',
                              x=340, y=450)

        self.label9 = MyLabel(self.frame1, text='Сидорик Валерий Владимирович', font_size=12, font_weight='bold', x=340,
                              y=480)

        self.label10 = MyLabel(self.frame1, text='Минск 2024', font_size=12, font_weight='bold', x=350, y=630)

        self.label11 = MyLabel(self.frame1, text='Учет дежурств в общежитии', font_size=16, font_weight='bold', x=235,
                               y=240)

        # Лейблы с текстом окна об авторе
        self.label12 = MyLabel(self.frameAuthor, text='Автор', font_size=8, x=170, y=305, bg='#e3e2e1', fg = "#616161")

        self.label14 = MyLabel(self.frameAuthor, text='Жоров Евгений Александрович', font_size=11, font_weight='bold', bg='#e3e2e1',
                               x=70, y=330)

        self.label15 = MyLabel(self.frameAuthor, text='vanasokolov844@gmail.com', font_size=9, font_weight='bold',
                               x=125, y=446)

        # Лейблы для текста окна о программе
        self.label16 = MyLabel(self.frameAbtProgramm, text='Учет дежурств в общежитии', font_size=18,
                               font_weight='bold', x=380, y=40)

        self.label17 = MyLabel(self.frameAbtProgramm, text='Программа позволяет:', font_size=13, font_weight='bold',
                               x=440, y=90)

        self.label18 = MyLabel(self.frameAbtProgramm, text='1. Записывать данные дежурного студента\n'
                                                           '2. Удалять записи на дежурства\n'
                                                           '3. Редактировать данные дежурного студента\n'
                                                           '4. Просматривать результат в главном окне\n'
                                                           '5. Сохранять записи в файл форматов .docx .xlsx\n'
                                                           '6. Считывать данные с Excel\n'
                                                           '7. Отправлять docx файл в Telegram-бот', font_size=11,
                               justify='left', bg='#CDDBFE', x=350, y=130)

        self.label19 = MyLabel(self.frameAbtProgramm, text='Версия: 1.0.0.2024', font_size=10, x=350, y=710)

        # Лейблы с текстом главного окна
        self.label20 = MyLabel(self.frameMain, text="Запись на дежурство", font_size=11, font_weight='bold', bg='#F1F1F1',
                               x=300, y=15)

        self.label21 = MyLabel(self.frameMain, text="Фамилия", font_size=10, font_weight='bold', x=20, y=105)
        self.label22 = MyLabel(self.frameMain, text="Имя", font_size=10, font_weight='bold', x=150, y=105)
        self.label23 = MyLabel(self.frameMain, text="Блок", font_size=10, font_weight='bold', x=450, y=105)
        self.label24 = MyLabel(self.frameMain, text="Номер телефона", font_size=10, font_weight='bold', x=280, y=105)
        self.label24 = MyLabel(self.frameMain, text="Время", font_size=10, font_weight='bold', x=400, y=175)
        self.label25 = MyLabel(self.frameMain, text="Месяц", font_size=10, font_weight='bold', x=150, y=175)
        self.label26 = MyLabel(self.frameMain, text="День", font_size=10, font_weight='bold', x=20, y=175)
        self.label27 = MyLabel(self.frameMain, text="Комната", font_size=10, font_weight='bold', x=540, y=105)
        self.label28 = MyLabel(self.frameMain, text="Год", font_size=10, font_weight='bold', x=280, y=175)

        self.label29 = MyLabel(self.frameAuthentification, text = "Добро пожаловать!", font_size = 18, font_weight='bold', x=290, y=260)

        self.label30 = MyLabel(self.frameAuthentification, text = "Войдите в программу как комендант/студент", font_size = 10, x = 270, y = 300)

        self.label31 = MyLabel(self.frameAuthentification, text = "Логин", font_size=10, x = 260, y = 360)

        self.label32 = MyLabel(self.frameAuthentification, text = "Пароль", font_size=10, x = 260, y = 430)

        self.label33 = MyLabel(self.frameAuthentification, text = 'Не комендант?',font_size=9, x = 300, y = 570)

        self.label34 = MyLabel(self.frameAuthentification, text = 'Войти как студент',font_size=9, font_weight='bold', x = 395, y = 570, cursor = "hand2", fg = '#3469ed')

        self.label34.bind("<Button-1>", lambda e: self.fromAuthentificToMainStudent())

        self.label35 = MyLabel(self.frameAuthor, text='Почта', font_size=7, fg = "#616161", x=125, y=435)

        self.label36 = MyLabel(self.frameAuthor, text='+375297379694', font_size=9, font_weight='bold', x=125, y=499)

        self.label37 = MyLabel(self.frameAuthor, text='Номер телефона', font_size=7, fg="#616161", x=125, y=486)

        self.label38 = MyLabel(self.frameAuthor, text='Беларусь, г. Минск', font_size=9, font_weight='bold', x=125, y=549)

        self.label39 = MyLabel(self.frameAuthor, text='Местоположение', font_size=7, fg="#616161", x=125, y=536)

        self.label40 = MyLabel(self.frameAuthor, text='ОБО МНЕ', font_size=17, font_weight='bold', x=385, y=150)

        self.label41 = MyLabel(self.frameAuthor, text='Привет! Меня зовут Женя.\nЯ студент БНТУ ФИТР гр.10701123', font_size=10, x=385, y=200)

        self.label42 = MyLabel(self.frameAuthor,
                               text='Программирование для меня не просто хобби, а настоящее увлечение.'
                                    ' Я получаю огромное удовольствие от создания новых проектов '
                                    'и решения сложных задач.'
                                    ' Мне нравится изучать новые языки программирования и технологии, а также '
                                    'применять их на практике. Я всегда стремлюсь к саморазвитию и стараюсь держать '
                                    'руку на пульсе новых IT-трендов.', font_size=10, wraplength=350, x=385,y=250)

        self.label43 = MyLabel(self.frameMain, text="Сортировка", font_size=8, font_weight='bold',fg="#545454", x=20, y=60)
        self.label43.bind('<Enter>', self.cursorOnLabelSort)
        self.label43.bind('<Leave>', self.cursorNotOnLabelSort)
        self.label43.bind('<Button-1>', self.toggleDropdown)

        self.label44 = MyLabel(self.frameMain, text="Поиск", font_size=8, font_weight='bold', fg="#545454", x=134, y=60)
        self.label44.bind('<Enter>', self.cursorOnLabelSearch)
        self.label44.bind('<Leave>', self.cursorNotOnLabelSearch)
        self.label44.bind('<Button-1>', self.toggleDropdownSearch)

        # Entry поля на главном окне
        # Ввод фамилии
        self.entry1 = tk.Entry(self.frameMain, width=15, borderwidth=2)
        self.entry1.place(x=20, y=135)

        # Ввод имени
        self.entry2 = tk.Entry(self.frameMain, width=15, borderwidth=2)
        self.entry2.place(x=150, y=135)

        # Ввод блока
        self.entry3 = tk.Entry(self.frameMain, width=10, borderwidth=2)
        self.entry3.place(x=450, y=135)

        # Ввод номера тф
        self.entry4 = tk.Entry(self.frameMain, width=18, borderwidth=2)
        self.entry4.place(x=280, y=135)

        #Ввод логина
        self.entry5 = tk.Entry(self.frameAuthentification, width=30, font=('Montserrat', 13))
        self.entry5.place(x = 260, y = 390)

        # Ввод пароля
        self.entry6 = tk.Entry(self.frameAuthentification, width=30, font=('Montserrat', 13), show = "*")
        self.entry6.place(x=260, y=460)

        self.dictForMonthDays = {"Январь": 31,
                                 "Февраль": 28,  # Простой год
                                 "Март": 31,
                                 "Апрель": 30,
                                 "Май": 31,
                                 "Июнь": 30,
                                 "Июль": 31,
                                 "Август": 31,
                                 "Сентябрь": 30,
                                 "Октябрь": 31,
                                 "Ноябрь": 30,
                                 "Декабрь": 31}

        # Ввод месяца
        self.comboboxMonth = ttk.Combobox(self.frameMain, values=['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль',
                                                  'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь'], width=12,
                                          state="readonly")
        self.comboboxMonth.place(x=150, y=205)
        self.comboboxMonth.bind("<<ComboboxSelected>>", self.UpdateDays)

        self.comboboxDay = ttk.Combobox(self.frameMain, width=12, state="readonly")
        self.comboboxDay.place(x=20, y=205)
        self.comboboxDay.bind("<<ComboboxSelected>>", self.getInfoComboboxDay)

        self.comboboxTimeValues = ['8.00 - 10.30', '10.30 - 13.00', '13.00 - 15.30', '15.30 - 18.00', '18.00 - 20.00']

        self.comboboxTime = ttk.Combobox(self.frameMain, values=self.comboboxTimeValues, width=11, state="readonly")
        self.comboboxTime.place(x=400, y=205)

        self.comboboxRoomInBlock = ttk.Combobox(self.frameMain, values=['А', 'Б'], width=11, state='readonly')
        self.comboboxRoomInBlock.place(x=540, y=135)

        self.comboboxYear = ttk.Combobox(self.frameMain, values=['2025'], width=11, state='readonly')
        self.comboboxYear.place(x=280, y=205)
        self.comboboxYear.bind("<<ComboboxSelected>>", self.onYearChange)

        self.frameSort = tk.Frame(self.frameMain, width=180, height=150, bg="#f1f1f1", borderwidth=3, relief='ridge')

        self.labelSort1 = MyLabel(self.frameSort, text = "Сортировать по", font_weight="bold", font_size=8, x = 10, y = 10)

        self.comboboxSort = ttk.Combobox(self.frameSort, values=['По фамилиям (от А до Я)', 'По фамилиям (от Я до А)', 'По возрастанию блоков',
                                                                  'По убыванию блоков', 'По дате'], font = ('Montserrat', 8), width=22, state='readonly')
        self.comboboxSort.place(x = 10, y = 35)



        self.frameSearch = tk.Frame(self.frameMain, width=350, height=350, bg = "#f1f1f1", borderwidth=3, relief='ridge')

        self.labelSearch1 = MyLabel(self.frameSearch, text = "Выполнить поиск по", font_weight="bold", font_size=8, x = 20, y = 15)

        self.labelSearch2 = MyLabel(self.frameSearch, text = "Ваш запрос",fg = "#545454", font_size=8, x = 20, y = 130)

        self.searchOption = tk.StringVar()

        self.radioLastName = ttk.Radiobutton(self.frameSearch, text="Фамилиям", variable=self.searchOption, value = "lastName", command=self.updateEntryField)
        self.radioLastName.place(x = 20, y = 50)
        self.radioLastName.bind('<FocusIn>', self.removeFocus)

        self.radioBlocks = ttk.Radiobutton(self.frameSearch, text="Блокам", variable=self.searchOption, value = "blocks", command=self.updateEntryField)
        self.radioBlocks.place(x=20, y=75)
        self.radioBlocks.bind('<FocusIn>', self.removeFocus)

        self.radioDates = ttk.Radiobutton(self.frameSearch, text="Дате", variable=self.searchOption, value="dates", command=self.updateEntryField)
        self.radioDates.place(x=20, y=100)
        self.radioDates.bind('<FocusIn>', self.removeFocus)

        self.entryLastNamesSearch = tk.Entry(self.frameSearch, width=25, font = ('Montserrat',9))
        self.entryLastNamesSearch.place(x = 20, y = 150)

        self.entryBlocksSearch = tk.Entry(self.frameSearch, width=25, font=('Montserrat', 9))
        self.entryBlocksSearch.place(x=20, y=150)
        self.entryBlocksSearch.place_forget()

        self.entryDatesSearch = tk.Entry(self.frameSearch, width=25, font=('Montserrat', 9))
        self.entryDatesSearch.place(x=20, y=150)
        self.entryDatesSearch.place_forget()

        self.textToSearch = tk.Text(self.frameSearch, width=35, height=9, state='disabled')
        self.textToSearch.place(x=20, y=178)

        # Создание вертикального скроллбара
        self.scrollbar = tk.Scrollbar(self.frameSearch, orient='vertical', command=self.textToSearch.yview)
        self.scrollbar.place(x=315, y=178, height=144)

        # Связывание скроллбара с текстовым полем
        self.textToSearch.config(yscrollcommand=self.scrollbar.set)

        self.UpdateDays(None)

        # ---------------------------------------------------------------------------------------КАРТИНКИ
        self.imageDormitory = tk.PhotoImage(file='images//icons//DormitorySplash.png')
        self.labelForDormitory = MyLabel(self.frame1, image=self.imageDormitory, x=160, y=370)
        self.labelForDormitory.image = self.imageDormitory

        self.imagePlus = tk.PhotoImage(file='images//icons//Sort.png')
        self.labelForPlus = MyLabel(self.frameMain, image=self.imagePlus, x=91, y=61)
        self.labelForPlus.image = self.imagePlus

        self.imageSearch = tk.PhotoImage(file='images//icons//Search.png')
        self.labelForSearch = MyLabel(self.frameMain, image=self.imageSearch, x=173, y=61)
        self.labelForSearch.image = self.imageSearch

        self.imageSort2 = tk.PhotoImage(file='images//icons//Sort2.png')
        self.labelSort2 = MyLabel(self.frameSort, image=self.imageSort2, x=108, y=10)
        self.labelSort2.image = self.imageSort2

        self.imageSearchAcc = tk.PhotoImage(file='images//icons//SearchAcc.png')
        self.labelSearchAcc = MyLabel(self.frameSearch, image=self.imageSearchAcc, x=139, y=14)
        self.labelSearchAcc.image = self.imageSearchAcc

        self.imageMe = tk.PhotoImage(file='images//me//Me2.png')
        self.resizedImageMe = self.imageMe.subsample(9)
        self.labelForMe = MyLabel(self.frameAuthor, image=self.resizedImageMe, x=60, y=30)
        self.labelForMe.image = self.resizedImageMe

        self.imageGmail = tk.PhotoImage(file='images//icons//Gmail.png')
        self.labelForGmail = MyLabel(self.frameAuthor, image=self.imageGmail, x=90, y=443)
        self.labelForGmail.image = self.imageGmail

        self.imageSmartphone = tk.PhotoImage(file='images//icons//Smartphone.png')
        self.labelForSmartphone = MyLabel(self.frameAuthor, image=self.imageSmartphone, x=90, y=490)
        self.labelForSmartphone.image = self.imageSmartphone

        self.imageLocation = tk.PhotoImage(file='images//icons//Location.png')
        self.labelForLocation = MyLabel(self.frameAuthor, image=self.imageLocation, x=90, y=540)
        self.labelForLocation.image = self.imageLocation

        self.imageForAbtProgramm1 = tk.PhotoImage(file='images//icons//DormitoryWhite.png')
        self.labelForAbtProgramm1 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm1, bg="#3469ED", x=34,y=236)
        self.labelForAbtProgramm1.image = self.imageForAbtProgramm1

        self.imageForAbtProgramm2 = tk.PhotoImage(file='images//icons//Room.png')
        self.labelForAbtProgramm2 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm2, bg='#3469ED', x=150,
                                            y=235)
        self.labelForAbtProgramm2.image = self.imageForAbtProgramm2

        self.imageForAbtProgramm3 = tk.PhotoImage(file='images//icons//File.png')
        self.labelForAbtProgramm3 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm3, bg='#3469ED', x=30,
                                            y=350)
        self.labelForAbtProgramm3.image = self.imageForAbtProgramm3

        self.imageForAbtProgramm4 = tk.PhotoImage(file='images//icons//Excel.png')
        self.labelForAbtProgramm4 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm4, bg='#3469ED', x=150,
                                            y=350)
        self.labelForAbtProgramm4.image = self.imageForAbtProgramm4

        self.imageForAbtProgramm5 = tk.PhotoImage(file='images//icons//Human2.png')
        self.labelForAbtProgramm5 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm5, bg='#3469ED', x=30,
                                            y=460)
        self.labelForAbtProgramm5.image = self.imageForAbtProgramm5

        self.imageForAbtProgramm6 = tk.PhotoImage(file='images//icons//Telephone.png')
        self.labelForAbtProgramm6 = MyLabel(self.frameAbtProgramm, image=self.imageForAbtProgramm6, bg='#3469ED', x=150,
                                            y=460)
        self.labelForAbtProgramm6.image = self.imageForAbtProgramm6

        self.imageForAuthentific = tk.PhotoImage(file = 'images//icons//Dormitory90.png')
        self.labelForAuthentific = MyLabel(self.frameAuthentification, image=self.imageForAuthentific, x=365, y=165)
        self.labelForAuthentific.image = self.imageForAuthentific

        self.imageForDormitory32 = tk.PhotoImage(file='images//icons//Dormitory32.png')
        self.labelForDormitory32 = MyLabel(self.frameMain, image=self.imageForDormitory32, x=395, y=705)
        self.labelForDormitory32.image = self.imageForDormitory32

        # Картинка для кнопки выйти
        self.imageExit = tk.PhotoImage(file='images//buttons//Exit.png')

        # Картинка для кнопки назад
        self.imageBack = tk.PhotoImage(file='images//buttons//Back.png')

        # Картинка для кнопки об авторе
        self.imageHuman = tk.PhotoImage(file='images//buttons//Human.png')

        # Картинка для кнопки о программе
        self.imageAbtProgramm = tk.PhotoImage(file='images//buttons//About.png')

        # Картинка для кнопки на главную
        self.imageToTheMain = tk.PhotoImage(file='images//buttons//Home.png')

        # Картинка для кнопки записаться
        self.imageWrite = tk.PhotoImage(file='images//buttons//Write.png')

        # Картинка для кнопки очистки
        self.imageClear = tk.PhotoImage(file='images//buttons//Clear.png')

        # Картинка для кнопки удалить данные
        self.imageDelete = tk.PhotoImage(file='images//buttons//Delete.png')

        # Картинка для кнопки выбора сортировки
        self.imageSort = tk.PhotoImage(file='images//buttons//Sort.png')

        # Картинка для кнопки записи в Excel
        self.imageExcelBtn = tk.PhotoImage(file = 'images//buttons//Excel.png')

        #Картинка для кнопки отменить изменения
        self.imageCancel = tk.PhotoImage(file = 'images//buttons//Cancel.png')

        #Картинка для кнопки сохранить изменения
        self.imageSaveChanges = tk.PhotoImage(file = 'images//buttons//Save.png')

        # Картинка для кнопки отправить коменданту
        self.imageWord = tk.PhotoImage(file='images//buttons//Word.png')

        # ---------------------------------------------------------------------------------КАРТИНКИ

        # ---------------------------------------------------------------------------------КНОПКИ

        # Кнопка далее
        self.buttonStart = MyButton(self.frame1, text='Далее', fg='white', width=30, height=3,
                                    font_size=9, font_weight='bold', bg='#3469ed', command=self.openAuthentificWindow, x=150,
                                    y=700)
        #Кнопка войти
        self.buttonLogIn = MyButton(self.frameAuthentification, text = "Войти", fg='white', width=38, height=2,
                                    font_size=9, font_weight='bold', bg='#3469ed', command=self.fromAuthentificToMainCommandant, x=260,
                                    y=515)

        # Кнопка выход
        if self.frameMain:
            self.buttonExit = MyButton(self.frameMain, text='Выход', width=100, height=30, image=self.imageExit,
                                       compound=tk.RIGHT, font_size=9, font_weight='bold', command=self.exitAppMain, x=525,
                                       y=760)

        if self.frame1:
            self.buttonExit = MyButton(self.frame1, text='Выход', width=200, height=45, image=self.imageExit,
                                       compound=tk.RIGHT, font_size=9, font_weight='bold', command=self.exitAppSplash, x=410,
                                       y=700)

        # Кнопка назад УНИВЕРСАЛЬНАЯ
        if self.frameMain:
            self.buttonBack = MyButton(self.frameMain, text='Назад', width=100, height=30, image=self.imageBack,
                                       compound=tk.RIGHT, font_size=9, font_weight='bold',
                                       command=self.backFromMainto1st, x=645, y=760)

        if self.frameAbtProgramm:
            self.buttonBack = MyButton(self.frameAbtProgramm, text='Назад', width=100, height=30, image=self.imageBack,
                                       compound=tk.RIGHT, font_size=9, font_weight='bold',
                                       command=self.backFromAbtProgrammToMain, x=645, y=700)

        # Кнопка об авторе
        self.buttonAuthor = MyButton(self.frameMain, text='Об авторе', width=100, height=30, image=self.imageHuman,
                                     compound=tk.RIGHT, font_size=9, font_weight='bold',
                                     command=self.openAuthorFromMain, x=200, y=760)

        # Кнопка о программе
        self.buttonAbtProgramm = MyButton(self.frameMain, text='О программе', width=120, height=30,
                                          image=self.imageAbtProgramm,
                                          compound=tk.RIGHT, font_size=9, font_weight='bold',
                                          command=self.openAbtProgramm, x=60, y=760)
        #Кнопка отсортировать
        self.buttonSort = MyButton(self.frameSort, text='Отсортировать',width=14, height=1, bg = "#3469ed",
                                   font_size=9, font_weight='bold', fg = "#f1f1f1", x=34, y=105, command=self.funcOfSort)

        # Кнопка поиска
        self.buttonSearch = MyButton(self.frameSearch, text='Поиск', width=12, height=1, bg="#3469ed",
                                   font_size=9, font_weight='bold', fg="#f1f1f1", x=210, y=145, command=self.performSearch)

        # Кнопка на главную
        self.buttonToTheMain = MyButton(self.frameAuthor, text='На главную', width=250, height=30,
                                        image=self.imageToTheMain,
                                        compound=tk.RIGHT, font_size=9, font_weight='bold',
                                        command=self.backFromAuthorToMain, x=385, y=400)
        # Кнопка записаться
        self.buttonZapisat = MyButton(self.frameMain, text='Записаться', width=80, height=60, image=self.imageWrite,
                                      compound=tk.TOP, font_size=9, font_weight='bold', command=self.GetResults, x=655,
                                      y=140)

        self.buttonSupportAuth = MyButton(self.frameAuthor, text="Поддержать автора", fg='white', width=25, height=2,
                                    font_size=9, font_weight='bold', bg='#3469ed', x=100, y=595)

        # Кнопка очистить ввод
        self.buttonClear = MyButton(self.frameMain, text="Очистить\nввод", width=80, height=60, image=self.imageClear,
                                    compound=tk.TOP, font_size=9, font_weight='bold', command=self.clearEntry, x=655,
                                    y=250)

        # Кнопка удалить данные
        self.buttonDeleteData = MyButton(self.frameMain, text='Удалить\nданные', width=80, height=60,
                                         image=self.imageDelete,
                                         compound=tk.TOP, font_size=9, font_weight='bold', command=self.deleteAllData,
                                         x=655, y=360)

        # Кнопка записи в Excel
        self.buttonToExcel = MyButton(self.frameMain, text='Сохранить\nв Excel', width=80, height=60,
                                         image=self.imageExcelBtn,
                                         compound=tk.TOP, font_size=9, font_weight='bold', command=self.infoToExcel,
                                         x=655, y=470)


        # Кнопка сохранения в Docx
        self.buttonToDocx = MyButton(self.frameMain, text='Сохранить\nв Word', width=80, height=60,
                                          image=self.imageWord,
                                          compound=tk.TOP, font_size=9, font_weight='bold', command = self.infoToDocx,
                                          x = 655, y = 580)

        # Кнопка отправки файла в бот
        self.buttonSendToBot = MyButton(self.frameMain, text='Отправить\nкоменданту', width=80, height=60,
                                     image=self.imageWord,
                                     compound=tk.TOP, font_size=9, font_weight='bold', command=self.sendFileToBot,
                                     x=655, y=630)

        #Кнопка сохранения изменений
        self.buttonSaveChanges = MyButton(self.frameMain, text = "Сохранить\nизменения", width=80, height=60,
                                         image=self.imageSaveChanges,
                                         compound=tk.TOP, font_size=9, font_weight='bold', command=self.saveAndPlace,
                                         x=655, y=100)
        self.buttonSaveChanges.place_forget()

        #Кнопка отмена изменений
        self.buttonCancel = MyButton(self.frameMain, text="Отмена", width=80, height=60,
                                          image=self.imageCancel,
                                          compound=tk.TOP, font_size=9, font_weight='bold', command=self.cancelChange,
                                          x=655, y=210)
        self.buttonCancel.place_forget()

        # Кнопка удаления записи
        self.buttonDeleteRecord = MyButton(self.frameMain, text='Удалить\nзапись', width=80, height=60, image=self.imageDelete,
                                      compound=tk.TOP, font_size=9, font_weight='bold', command=self.deleteRecord, x=655,
                                      y=700)
        self.buttonDeleteRecord.place_forget()

        #--------------------------------------------------------------------------------------------------------------МЕНЮ КНОПКИ

        self.buttonMenuFile = tk.Menubutton(self.frameMain, text = 'Файл', relief=tk.RAISED, width = 5)
        self.buttonMenuFile.place(x = 0, y = 0)

        self.fileMenu = tk.Menu(self.buttonMenuFile, tearoff=0)
        self.buttonMenuFile.config(menu = self.fileMenu)

        #Открытие файла (чтение из файла)
        self.fileMenu.add_command(label = "Открыть",command=self.openExcelFile)

        self.fileMenu.add_command(label = "Очистить файл", command = self.clearDataFromExcel)

        #Помощь
        self.buttonMenuHelp = tk.Menubutton(self.frameMain, text = "Помощь", relief = tk.RAISED, width = 7)
        self.buttonMenuHelp.place(x = 43, y = 0)

        self.helpMenu = tk.Menu(self.buttonMenuHelp, tearoff=0)
        self.buttonMenuHelp.config(menu = self.helpMenu)

        self.helpMenu.add_command(label = "Использование", command = self.openHelpWindow)

        # --------------------------------------------------------------------------------------------------------------МЕНЮ КНОПКИ

    # -----------------------------------------------------------------------------------------------------------------КНОПКИ

    # -----------------------------------------------------------------------------------------------------------------ФУНКЦИИ

    def onClosing(self):
        self.telegramBot.bot.stop_polling()  # Останавливаем бота при закрытии окна
        self.destroy() # Закрываем главное окно

    #Функция задающая фокус при выборе радиокнопки
    def removeFocus(self, event):
        self.frameSearch.focus()

    # Закрытие окна
    def exitAppMain(self):
        result = tk.messagebox.askyesnocancel("Выход", "Все несохраненные данные будут удалены.\nВыйти?")

        if result:

            self.destroy()

    def exitAppSplash(self):
        self.destroy()

    def checkInternetConnection(self, url='http://www.google.com/', timeout=5):
        try:
            requests.get(url, timeout=timeout)
            return True
        except requests.ConnectionError:
            return False

    def cursorOnLabelSort(self, event):

        self.label43.config(fg = "black", cursor = 'hand2')

    def cursorNotOnLabelSort(self, event):

        self.label43.config(fg = "#545454")

    def cursorOnLabelSearch(self, event):

        self.label44.config(fg = "black", cursor = 'hand2')

    def cursorNotOnLabelSearch(self, event):

        self.label44.config(fg = "#545454")

    def toggleDropdown(self, event):
        if self.dropDownVisibleSort:

            self.comboboxSort.config(state='normal')
            self.comboboxSort.delete(0, tk.END)
            self.comboboxSort.config(state='readonly')

            self.frameSort.place_forget()
        else:
            self.frameSort.place(x=20, y=85)
        self.dropDownVisibleSort = not self.dropDownVisibleSort

    #Функция закрывающая фреймы
    def closeFrame(self, event):

        if self.frameSort:

            self.comboboxSort.config(state = 'normal')
            self.comboboxSort.delete(0, tk.END)
            self.comboboxSort.config(state = 'readonly')

            self.frameSort.place_forget()

            self.dropDownVisibleSort = False

        if self.frameSearch:

            self.clearSelectionRadio()
            self.textToSearch.config(state='normal')
            self.textToSearch.delete("1.0", tk.END)
            self.textToSearch.config(state='disabled')
            self.frameSearch.place_forget()
            self.dropDownVisibleSearch = False

    def toggleDropdownSearch(self, event):
        if self.dropDownVisibleSearch:

            self.frameSearch.place_forget()
        else:
            self.frameSearch.place(x=150, y=85)
        self.dropDownVisibleSearch = not self.dropDownVisibleSearch

    #Функция снимающая метку с радиокнопки
    def clearSelectionRadio(self):

        searchType = self.searchOption.get()

        if searchType:
            self.labelSearch2.config(text="Ваш запрос")

        self.searchOption.set("")

    def updateEntryField(self):
        searchType = self.searchOption.get()
        # Скрываем все поля ввода сначала
        self.entryLastNamesSearch.place_forget()
        self.entryBlocksSearch.place_forget()
        self.entryDatesSearch.place_forget()

        # Отображаем соответствующее поле ввода в зависимости от выбранной радиокнопки
        if searchType == "lastName":
            self.entryLastNamesSearch.place(x=20, y=150)
            self.labelSearch2.config(text="Ваш запрос (Фамилия)")

        elif searchType == "blocks":
            self.entryBlocksSearch.place(x=20, y=150)
            self.labelSearch2.config(text="Ваш запрос (Блок)")

        elif searchType == "dates":
            self.entryDatesSearch.place(x=20, y=150)
            self.labelSearch2.config(text="Ваш запрос (XX.YY.2025)")

    #Функция алгоритма поиска
    def performSearch(self):

        searchType = self.searchOption.get()
        searchQuery = ""

        if searchType == "lastName":
            searchQuery = self.entryLastNamesSearch.get().strip()



        elif searchType == "blocks":
            searchQuery = self.entryBlocksSearch.get().strip()



        elif searchType == "dates":
            searchQuery = self.entryDatesSearch.get().strip()



        else:
            messagebox.showinfo("Ошибка", "Выберите тип поиска")
            return

        self.textToSearch.config(state='normal')
        self.textToSearch.delete(1.0, tk.END)

        if not searchQuery:
            messagebox.showinfo("Ошибка", "Введите поисковый запрос")
            self.textToSearch.config(state='disabled')
            return

        found = False  # Флаг для отслеживания наличия совпадений

        if searchType == "lastName":
            for i, second_name in enumerate(self.recordSecondNames):
                if searchQuery.lower() == second_name.lower():
                    found = True
                    record_info = f"Фамилия: {self.recordSecondNames[i]}\n" \
                                  f"Имя: {self.recordFirstNames[i]}\n" \
                                  f"Блок: {self.recordBlocks[i]}\n" \
                                  f"Комната: {self.recordRooms[i]}\n" \
                                  f"Номер телефона: {self.recordTelNumber[i]}\n" \
                                  f"Дата: {self.recordDates[i]}\n" \
                                  f"Время: {self.recordTimes[i]}\n\n"
                    self.textToSearch.insert(tk.END, record_info)

            self.entryLastNamesSearch.delete(0, tk.END)

        elif searchType == "blocks":
            for i, block in enumerate(self.recordBlocks):
                if searchQuery.lower() == block.lower():
                    found = True
                    record_info = f"Фамилия: {self.recordSecondNames[i]}\n" \
                                  f"Имя: {self.recordFirstNames[i]}\n" \
                                  f"Блок: {self.recordBlocks[i]}\n" \
                                  f"Комната: {self.recordRooms[i]}\n" \
                                  f"Номер телефона: {self.recordTelNumber[i]}\n" \
                                  f"Дата: {self.recordDates[i]}\n" \
                                  f"Время: {self.recordTimes[i]}\n\n"
                    self.textToSearch.insert(tk.END, record_info)

            self.entryBlocksSearch.delete(0, tk.END)

        elif searchType == "dates":
            for i, date in enumerate(self.recordDates):
                if searchQuery == date:
                    found = True
                    record_info = f"Фамилия: {self.recordSecondNames[i]}\n" \
                                  f"Имя: {self.recordFirstNames[i]}\n" \
                                  f"Блок: {self.recordBlocks[i]}\n" \
                                  f"Комната: {self.recordRooms[i]}\n" \
                                  f"Номер телефона: {self.recordTelNumber[i]}\n" \
                                  f"Дата: {self.recordDates[i]}\n" \
                                  f"Время: {self.recordTimes[i]}\n\n"
                    self.textToSearch.insert(tk.END, record_info)
            self.entryDatesSearch.delete(0, tk.END)

        if not found:
            messagebox.showinfo("Результат поиска", "Совпадений не найдено")

        self.textToSearch.config(state='disabled')

    #Проверка правильности ввода логина и пароля
    def checkLoginPassword(self):

        self.userLogin = self.entry5.get()
        self.userPassword = self.entry6.get()

        if(str(self.userLogin) == "commandant" and str(self.userPassword) == "12345678"): #Селфы можно убрать
            self.user = User(self.userLogin, self.roles["commandant"])


            return True

        else:
            tk.messagebox.showerror("Ошибка", "Неверный логин или пароль")
            self.entry5.delete(0, tk.END)
            self.entry6.delete(0, tk.END)

            return False

    # Открытие второго окна (забиндить и добавить эл-ты второго окна в фрейм 2)
    def openAuthentificWindow(self):
        self.frame1.pack_forget()
        self.frameAuthentification.pack()
        self.inactivity.stopCheck()

    def fromSplashToMain(self):
        self.frame1.pack_forget()
        self.frameMain.pack()
        self.inactivity.stopCheck()

    #Существует ли пользователь (для удалении окна со входом)
    def checkIsUserExists(self):

        if self.user is not None:

            self.buttonStart.config(command = self.fromSplashToMain)

    #Открытие главного окна после входа
    def fromAuthentificToMainCommandant(self):

        if (self.checkLoginPassword()):

            self.checkRoleAndDisplayButtons()

            self.frameAuthentification.pack_forget()
            self.frameMain.pack()

            self.checkIsUserExists()

        else:
            return

    #Вход как студент
    def fromAuthentificToMainStudent(self):

        self.user = User("student", "студент")

        self.frameAuthentification.pack_forget()
        self.frameMain.pack()

        self.checkRoleAndDisplayButtons()

        self.checkIsUserExists()

    # Возвращение со второго окна на первое(с основного в Splash)
    def backFromMainto1st(self):
        self.frameMain.pack_forget()
        self.frame1.pack()
        self.inactivity.startCheck()

    # Открытие окна об авторе с главного окна
    def openAuthorFromMain(self):
        self.frameMain.pack_forget()
        self.frameAuthor.pack()

    # Возвращение с окна об авторе на главную
    def backFromAuthorToMain(self):
        self.frameAuthor.pack_forget()
        self.frameMain.pack()

    def openAbtProgramm(self):
        self.frameMain.pack_forget()
        self.frameAbtProgramm.pack()

    def backFromAbtProgrammToMain(self):
        self.frameAbtProgramm.pack_forget()
        self.frameMain.pack()

    def openHelpWindow(self):

        self.helpWindow = HelpWindow(self.frameMain, self)

    #Функция убирающая кнопки в зависимости от роли пользователя
    def checkRoleAndDisplayButtons(self):

        if self.user.role == "студент":

            self.buttonDeleteData.place_forget()
            self.buttonToDocx.place_forget()
            self.buttonClear.place_forget()
            self.fileMenu.delete(1)

            self.buttonToExcel.place(x=655, y=400)
            self.buttonSendToBot.place(x = 655, y = 505)
            self.buttonClear.place(x = 655, y = 290)

            self.startReading()

        else:

            self.buttonSendToBot.place_forget()

    # Очистка всех полей при нажатии кнопки очистить ввод
    def clearEntry(self):

        self.comboboxMonth.config(state="normal")
        self.comboboxDay.config(state="normal")
        self.comboboxTime.config(state="normal")
        self.comboboxRoomInBlock.config(state="normal")
        self.comboboxYear.config(state="normal")

        self.entry1.delete(0, tk.END)
        self.entry2.delete(0, tk.END)
        self.entry3.delete(0, tk.END)
        self.entry4.delete(0, tk.END)
        self.comboboxMonth.delete(0, tk.END)
        self.comboboxDay.delete(0, tk.END)
        self.comboboxTime.delete(0, tk.END)
        self.comboboxRoomInBlock.delete(0, tk.END)
        self.comboboxYear.delete(0, tk.END)

        self.comboboxMonth.config(state="readonly")
        self.comboboxDay.config(state="readonly")
        self.comboboxTime.config(state="readonly")
        self.comboboxRoomInBlock.config(state="readonly")
        self.comboboxYear.config(state="readonly")

    def GetResults(self):

        self.getUserInput()

        if not self.checkAllFields():
            return

        if not self.validateInput():
            return

        self.processDateAndRecord()

        self.displayResults()

        self.clearEntry()
        self.recordData()

    def getUserInput(self):
        self.lastName = self.entry1.get()
        self.firstName = self.entry2.get()
        self.Block = self.entry3.get()
        self.TelNumber = self.entry4.get()
        self.Month = self.comboboxMonth.get()
        self.Day = self.comboboxDay.get()
        self.Time = self.comboboxTime.get()
        self.Room = self.comboboxRoomInBlock.get()
        self.Year = self.comboboxYear.get()

    def validateInput(self):
        if not self.validateLastName():
            return False
        if not self.validateFirstName():
            return False
        if not self.validateBlock():
            return False
        if not self.validateTelNumber():
            return False
        return True

    def validateLastName(self):
        for symbol in self.lastName:
            if symbol.isdigit() or symbol in "_-+=!@#$%*^()&?~/.,":
                tk.messagebox.showerror("Некорректный ввод", "Фамилия введена некорректно!")
                self.entry1.delete(0, tk.END)
                return False
        return True

    def validateFirstName(self):
        for symbol in self.firstName:
            if symbol.isdigit() or symbol in "_-+=!@#$%*^()&?~/., ":
                tk.messagebox.showerror("Некорректный ввод", "Имя введено некорректно!")
                self.entry2.delete(0, tk.END)
                return False
        return True

    def validateBlock(self):
        # Проверяем, что блок состоит из трех символов
        if len(self.Block) != 3:
            tk.messagebox.showerror("Некорректный ввод", "Блок должен состоять из трех цифр!")
            self.entry3.delete(0, tk.END)
            return False

        # Проверяем, что блок не начинается с нуля
        if self.Block[0] == '0':
            tk.messagebox.showerror("Некорректный ввод", "Блок не должен начинаться с нуля!")
            self.entry3.delete(0, tk.END)
            return False

        # Проверяем, что все символы в блоке являются цифрами
        for symbol in self.Block:
            if not symbol.isdigit():
                tk.messagebox.showerror("Некорректный ввод",
                                        "Блок введен некорректно! Все символы должны быть цифрами.")
                self.entry3.delete(0, tk.END)
                return False

        return True

    def checkPhoneNumber(self, phone_number):
        # Регулярное выражение для проверки номера телефона
        pattern = r"^\+375(29|33)\d{7}$"
        return re.match(pattern, phone_number) is not None

    def validateTelNumber(self):
        if not self.checkPhoneNumber(self.TelNumber):
            tk.messagebox.showerror("Некорректный ввод",
                                    "Номер введен некорректно! Номер должен начинаться с +37529 или +37533")
            self.entry4.delete(0, tk.END)
            return False
        return True

    def checkAllFields(self):

        self.textPlace.config(state='normal')

        if not all([self.lastName, self.firstName, self.Block, self.TelNumber, self.Month, self.Day, self.Time,
                    self.Room, self.Year]):
            tk.messagebox.showerror("Ошибка", "Введите все данные!")
            return False
        return True

    def processDateAndRecord(self):

        self.monthForSort = 0
        for nameOfMonth, MonthNumber in self.DictForMonthes.items():

            if self.Month == nameOfMonth:

                if int(self.Day) < 10:
                    self.concatenateString = "0" + self.Day + "." + MonthNumber + "." + self.Year

                else:
                    self.concatenateString = self.Day + "." + MonthNumber + "." + self.Year

                self.monthForSort = int(MonthNumber)

                self.recordDuty(self.concatenateString, self.Time)

    def displayResults(self):

        self.fullStringSort = (f"{(str(self.lastName)).capitalize()}\t\t  {(str(self.firstName)).capitalize()}\t\t"
                               f"{str(self.Block) + str(self.Room)}\t\t{str(self.TelNumber)}\t\t   {str(self.concatenateString)}\t\t{str(self.Time)}\n")
        self.textPlace.insert(tk.END,
                              f"{(str(self.lastName)).capitalize()}\t\t  {(str(self.firstName)).capitalize()}\t\t"
                              f"{str(self.Block) + str(self.Room)}\t\t{str(self.TelNumber)}\t\t   {str(self.concatenateString)}\t\t{str(self.Time)}\n")
        self.textPlace.config(state='disabled')

    def onYearChange(self, event):
        self.updateAvailableTimes()

    def UpdateDays(self, event):

        self.monthUser = self.comboboxMonth.get()
        current_day = self.comboboxDay.get()

        self.comboboxTime.config(state = 'normal')
        self.comboboxTime.delete(0, tk.END)
        self.comboboxTime.config(state='readonly')

        # Изначально устанавливаем 31 день
        self.days = [str(i) for i in range(1, 32)]

        if self.monthUser:
            days_in_month = self.dictForMonthDays[self.monthUser]
            self.days = [str(i) for i in range(1, days_in_month + 1)]

            if current_day and int(current_day) > days_in_month:
                self.comboboxDay.set("")
            elif current_day:
                self.comboboxDay.set(current_day)

        self.comboboxDay['values'] = self.days
        self.updateAvailableTimes()

    def getInfoComboboxDay(self, event):
        self.updateAvailableTimes()

    def updateAvailableTimes(self):

        stringForDay = self.comboboxDay.get()
        stringForMonth = self.comboboxMonth.get()
        stringForYear = self.comboboxYear.get()

        self.comboboxTime.config(state='normal')
        self.comboboxTime.delete(0, tk.END)
        self.comboboxTime.config(state='readonly')

        if not stringForDay or not stringForMonth or not stringForYear:  # Проверка на пустую строку
            self.comboboxTime['values'] = self.comboboxTimeValues
            return

        self.FullConcatenate = None
        for nameOfMonth, MonthNumber in self.DictForMonthes.items():
            if stringForMonth == nameOfMonth:
                if int(stringForDay) < 10:
                    self.FullConcatenate = "0" + stringForDay + "." + MonthNumber + "." + stringForYear
                else:
                    self.FullConcatenate = stringForDay + "." + MonthNumber + "." + stringForYear

        if self.FullConcatenate:
            if self.FullConcatenate in self.dictZapisanye:
                reserved_times = self.dictZapisanye[self.FullConcatenate]
                available_times = [time for time in self.comboboxTimeValues if time not in reserved_times]
                self.comboboxTime['values'] = available_times
            else:
                self.comboboxTime['values'] = self.comboboxTimeValues

    def recordDuty(self, date, time):
        if date in self.dictZapisanye:
            if time not in self.dictZapisanye[date]:
                self.dictZapisanye[date].append(time)
        else:
            self.dictZapisanye[date] = [time]

    def recordData(self):

        self.recordBlocks.append(self.Block)
        self.recordSecondNames.append(self.lastName.capitalize())
        self.recordFullStrings.append(self.fullStringSort)
        self.recordDays.append(int(self.Day))
        self.recordMonthes.append(self.monthForSort)
        self.recordRooms.append(self.Room)
        self.recordFirstNames.append(self.firstName.capitalize())
        self.recordTelNumber.append(self.TelNumber)
        self.recordDates.append(self.concatenateString)
        self.recordTimes.append(self.Time)
        self.recordYears.append(self.Year)

    # Удаление всех данных
    def deleteAllData(self):

        self.textPlace.config(state='normal')

        result = tk.messagebox.askyesno('Удаление данных',
                                        'Все данные о дежурстве будут безвозвратно удалены.\nПродолжить?')

        if result:
            self.textPlace.delete('3.0', tk.END)
            self.textPlace.insert('2.0', '\n')

        self.textPlace.config(state='disabled')

        self.recordBlocks.clear()
        self.recordSecondNames.clear()
        self.recordFullStrings.clear()
        self.recordTelNumber.clear()
        self.recordRooms.clear()
        self.recordFirstNames.clear()
        self.recordDates.clear()
        self.recordTimes.clear()

        self.dictZapisanye.clear()

    #Функция отмены изменений
    def cancelChange(self):
        self.clearEntry()

        # Скрытие кнопок сохранения изменений и отмены
        self.buttonSaveChanges.place_forget()
        self.buttonCancel.place_forget()
        self.buttonDeleteRecord.place_forget()

        # Возвращение кнопок на исходные позиции
        self.buttonZapisat.place(x=645, y=100)
        self.buttonToExcel.place(x=645, y=530)
        self.buttonDeleteData.place(x=645, y=320)
        self.buttonClear.place(x=645, y=210)
        self.buttonToDocx.place(x = 645, y = 630)

        # Удаление стрелочки и подсветки с текущей выбранной строки
        if self.selected_index is not None:
            self.removeArrowAndHighlight(self.selected_index)
            self.selected_index = None

    #Функция убирающая кнопки при нажатии на редактирование строки
    def switchButtons(self):

        self.buttonZapisat.place_forget()
        self.buttonToExcel.place_forget()
        self.buttonDeleteData.place_forget()
        self.buttonClear.place_forget()
        self.buttonToDocx.place_forget()
        self.buttonSendToBot.place_forget()

        self.buttonSaveChanges.place(x = 655, y = 100)
        self.buttonCancel.place(x = 655, y = 210)

        if self.user.role == "комендант":
            self.buttonDeleteRecord.place(x = 655, y = 320)

    def deleteRecord(self):
        if hasattr(self, 'selected_index') and self.selected_index is not None:
            index = self.selected_index
            row_index = int(index.split('.')[0]) - 3  # Корректировка для учета первых двух строк

            # Проверка роли пользователя
            if self.user.role == "комендант":
                # Получение старого времени и даты для удаления из словаря
                old_date_str = self.recordDates[row_index]
                old_time = self.recordTimes[row_index]
                self.removeDuty(old_date_str, old_time)

                # Определение, является ли удаляемая запись последней
                is_last_record = (row_index == len(self.recordFullStrings) - 1)

                if not is_last_record:
                    # Перемещение последней строки на место удаленной
                    self.recordSecondNames[row_index] = self.recordSecondNames[-1]
                    self.recordFirstNames[row_index] = self.recordFirstNames[-1]
                    self.recordBlocks[row_index] = self.recordBlocks[-1]
                    self.recordRooms[row_index] = self.recordRooms[-1]
                    self.recordTelNumber[row_index] = self.recordTelNumber[-1]
                    self.recordDays[row_index] = self.recordDays[-1]
                    self.recordMonthes[row_index] = self.recordMonthes[-1]
                    self.recordYears[row_index] = self.recordYears[-1]
                    self.recordDates[row_index] = self.recordDates[-1]
                    self.recordTimes[row_index] = self.recordTimes[-1]
                    self.recordFullStrings[row_index] = self.recordFullStrings[-1]

                # Удаление последней записи
                self.recordSecondNames.pop()
                self.recordFirstNames.pop()
                self.recordBlocks.pop()
                self.recordRooms.pop()
                self.recordTelNumber.pop()
                self.recordDays.pop()
                self.recordMonthes.pop()
                self.recordYears.pop()
                self.recordDates.pop()
                self.recordTimes.pop()
                self.recordFullStrings.pop()

                # Удаление строки из textPlace
                self.textPlace.config(state='normal')
                self.textPlace.delete("%s linestart" % index, "%s lineend" % index)

                # Перезапись всех строк для обновления отображения, сохраняя первую пустую строку
                self.textPlace.delete('3.0', tk.END)
                self.textPlace.insert('2.0', "\n")
                for i in range(len(self.recordFullStrings)):
                    self.textPlace.insert(f"{i + 3}.0", self.recordFullStrings[i] + "\n")

                self.textPlace.config(state='disabled')

                # Очистка выбранного индекса
                self.selected_index = None
                messagebox.showinfo("Удаление записи", "Запись успешно удалена.")

                self.cancelChange()

            else:
                messagebox.showwarning("Доступ запрещен", "У вас нет прав для удаления записей.")
        else:
            messagebox.showwarning("Удаление записи", "Пожалуйста, выберите запись для удаления.")

    #Функция получения данных при нажатии на текст
    def onTextClick(self, event):
        # Сохраняем индекс предыдущей выбранной строки
        if hasattr(self, 'selected_index') and self.selected_index is not None:
            previous_index = self.selected_index
        else:
            previous_index = None

        self.textPlace.config(state='normal')
        index = self.textPlace.index("@%s,%s" % (event.x, event.y))
        line_content = self.textPlace.get("%s linestart" % index, "%s lineend" % index).strip()
        self.textPlace.config(state='disabled')

        parts = line_content.split()
        if len(parts) >= 8:  # Обновлено для учета года
            try:
                row_index = int(index.split('.')[0]) - 3  # Корректировка для учета первых двух строк

                if row_index >= 0 and row_index < len(self.recordSecondNames):
                    lastName = self.recordSecondNames[row_index].strip()
                    firstName = self.recordFirstNames[row_index].strip()
                    block = self.recordBlocks[row_index].strip()
                    room = self.recordRooms[row_index].strip()
                    telNumber = self.recordTelNumber[row_index].strip()
                    day = str(self.recordDays[row_index]) if row_index < len(
                        self.recordDays) else '01'  # Проверка на наличие дня
                    month = self.recordMonthes[row_index] if row_index < len(
                        self.recordMonthes) else 1  # Проверка на наличие месяца
                    year = self.recordYears[row_index] if row_index < len(
                        self.recordYears) else '2023'  # Проверка на наличие года
                    date_str = self.recordDates[row_index].strip()
                    time = self.recordTimes[row_index].strip()

                    # Форматируем месяц как строку с двумя цифрами
                    month_str = f"{month:02}"

                    if self.user.role == "студент":
                        formatted_record = f"{lastName.capitalize()}\t\t  {firstName.capitalize()}\t\t{block}{room}\t\t{telNumber}\t\t   {day}.{month_str}.{year}\t\t{time}\n"
                        if formatted_record not in self.studentRecords:
                            tk.messagebox.showerror("Редактирование", "Ваша роль не позволяет изменять записи")
                            return

                    self.comboboxDay.set(day)
                    month_name = next((key for key, value in self.DictForMonthes.items() if value == month_str), None)
                    if month_name:
                        self.comboboxMonth.set(month_name)

                    self.comboboxYear.set(year)
                    self.comboboxTime.set(time)
                    self.updateAvailableTimes()

                    self.entry1.config(state='normal')
                    self.entry1.delete(0, tk.END)
                    self.entry1.insert(0, lastName)

                    self.entry2.config(state='normal')
                    self.entry2.delete(0, tk.END)
                    self.entry2.insert(0, firstName)

                    self.entry3.config(state='normal')
                    self.entry3.delete(0, tk.END)
                    self.entry3.insert(0, block)

                    self.comboboxRoomInBlock.config(state='normal')
                    self.comboboxRoomInBlock.set(room)

                    self.entry4.config(state='normal')
                    self.entry4.delete(0, tk.END)
                    self.entry4.insert(0, telNumber)

                    # Удаление стрелочки и подсветки с предыдущей строки
                    if previous_index is not None and previous_index != index:
                        previous_line_content = self.textPlace.get("%s linestart" % previous_index,
                                                                   "%s lineend" % previous_index).strip()
                        if previous_line_content.endswith("←"):
                            updated_previous_line = previous_line_content[:-1].rstrip()
                            self.textPlace.config(state='normal')
                            self.textPlace.delete("%s linestart" % previous_index, "%s lineend" % previous_index)
                            self.textPlace.insert("%s linestart" % previous_index, updated_previous_line)
                            self.textPlace.tag_remove("highlight", "%s linestart" % previous_index,
                                                      "%s lineend" % previous_index)
                            self.textPlace.config(state='disabled')

                    # Добавление стрелочки "←" к новой выбранной строке и подсветка
                    if not line_content.endswith("←"):
                        updated_line = line_content + " ←"
                        self.textPlace.config(state='normal')
                        self.textPlace.delete("%s linestart" % index, "%s lineend" % index)
                        self.textPlace.insert("%s linestart" % index, updated_line)
                    self.textPlace.tag_add("highlight", "%s linestart" % index, "%s lineend" % index)
                    self.textPlace.tag_configure("highlight", background="#89a8f5")
                    self.textPlace.config(state='disabled')

                    # Обновление индекса выбранной строки
                    self.selected_index = index
                    self.switchButtons()
                else:
                    print("Индекс строки выходит за пределы списка")
            except ValueError:
                print("Не удалось преобразовать индекс строки в целое число")

    def saveChanges(self):
        # Получаем обновленные данные из полей
        self.getUserInput()

        # Проверка все ли поля введены
        if not self.checkAllFields():
            self.removeArrowAndHighlight(self.selected_index)
            return

        # Проверка корректности ввода полей
        if not self.validateInput():
            self.removeArrowAndHighlight(self.selected_index)
            return

        # Обновление записи
        self.processDateAndRecord()

        self.displayResults()

        self.clearEntry()
        self.recordData()

    def removeDuty(self, date, time):
        if date in self.dictZapisanye:
            if time in self.dictZapisanye[date]:
                self.dictZapisanye[date].remove(time)
                if not self.dictZapisanye[date]:  # Удаление даты, если нет занятых временных слотов
                    del self.dictZapisanye[date]

        # Функция убирающая подсветку текста и стрелочку

    def removeArrowAndHighlight(self, index):
        if index is not None:
            line_content = self.textPlace.get("%s linestart" % index, "%s lineend" % index).strip()
            if line_content.endswith("←"):
                updated_line = line_content[:-1].rstrip()
                self.textPlace.config(state='normal')
                self.textPlace.delete("%s linestart" % index, "%s lineend" % index)
                self.textPlace.insert("%s linestart" % index, updated_line)
                self.textPlace.tag_remove("highlight", "%s linestart" % index, "%s lineend" % index)

    #Функция вызывающая функцию сохранения изменений и позиционирующая кнопки
    def saveAndPlace(self):

        self.saveChanges()
        self.cancelChange()

    #Функция записывающая данные в Excel
    def infoToExcel(self):

        result = tk.messagebox.askyesno("Изменить путь к файлу", "Хотите ли вы выбрать другой файл для записи?")

        if result:

            filePath = tk.filedialog.askopenfilename(title="Выберите файл", filetypes = [("Excel файлы", "*.xls"), ("Excel файлы", ".xlsx")])

            if filePath:

                self.excelObject = WorkExcel(self, filePath)

                self.excelObject.infoInFile(filePath)

                tk.messagebox.showinfo("Сохранение", f"Данные успешно сохранены!\n{filePath}")

        if not result:

            filePath = "excelFiles//Book1.xlsx"

            if filePath:

                self.excelObject = WorkExcel(self, filePath)

                self.excelObject.infoInFile(filePath)

                tk.messagebox.showinfo("Сохранение", f"Данные успешно сохранены!\n{filePath}")

    #Функция очистки файла по умолчанию
    def clearDataFromExcel(self):

        result = tk.messagebox.askyesno("Очистить файл", "Все данные из Excel файла по умолчанию будут удалены. Продолжить?")

        if result:

            filePath = "excelFiles//Book1.xlsx"

            self.excelObject = WorkExcel(self, filePath)

        else:

            return

    #Функция начального считывания файла Excel если пользователь студент
    def startReading(self):

        filePath = "excelFiles//Book1.xlsx"
        # Проверка, выбрал ли пользователь файл или нажал "Отмена"
        if not filePath:
            return

        try:
            self.wb = load_workbook(filePath)
            self.activeList = self.wb.active
        except Exception as e:
            tk.messagebox.showerror("Ошибка файла", f"Не удалось открыть файл: {str(e)}")
            return

        self.readFromExcel()

    # Считывание данных с Excel
    def readFromExcel(self):

        expectedHeaders = ['Фамилия', 'Имя', 'Блок', 'Комната', 'Номер', 'Дата', 'Время']
        actualHeaders = [self.activeList.cell(row=1, column=i + 1).value for i in range(len(expectedHeaders))]

        if actualHeaders != expectedHeaders:
            tk.messagebox.showerror("Ошибка файла", "Вы открыли файл с посторонними данными")
            return

        self.recordBlocks.clear()
        self.recordSecondNames.clear()
        self.recordFullStrings.clear()
        self.recordTelNumber.clear()
        self.recordRooms.clear()
        self.recordFirstNames.clear()
        self.recordDates.clear()
        self.recordTimes.clear()
        self.recordDays.clear()  # Очистка списка recordDays
        self.recordMonthes.clear()  # Очистка списка recordMonthes
        self.recordYears.clear()  # Очистка списка recordYears
        self.dictZapisanye.clear()

        for row in self.activeList.iter_rows(min_row=2, max_row=self.activeList.max_row, min_col=1,
                                             max_col=self.activeList.max_column, values_only=True):
            self.recordSecondNames.append(row[0])
            self.recordFirstNames.append(row[1])
            self.recordBlocks.append(row[2])
            self.recordRooms.append(row[3])
            self.recordTelNumber.append(row[4])
            self.recordDates.append(row[5])
            self.recordTimes.append(row[6])

            # Добавим извлечение дня, месяца и года из даты и занесем в соответствующие списки
            day, month, year = map(int, row[5].split('.'))
            self.recordDays.append(day)
            self.recordMonthes.append(month)
            self.recordYears.append(year)

            # Формирование полной строки и добавление в recordFullStrings
            full_string = f"{row[0].capitalize()}\t\t  {row[1].capitalize()}\t\t{row[2]}{row[3]}\t\t{row[4]}\t\t   {row[5]}\t\t{row[6]}"
            self.recordFullStrings.append(full_string)

            # Обновление словаря записанных временных слотов
            if row[5] not in self.dictZapisanye:
                self.dictZapisanye[row[5]] = []
            self.dictZapisanye[row[5]].append(row[6])

        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', '\n')

        for i in range(len(self.recordBlocks)):
            self.textPlace.insert(tk.END,
                                  f"{(str(self.recordSecondNames[i])).capitalize()}\t\t  {(str(self.recordFirstNames[i])).capitalize()}\t\t"
                                  f"{str(self.recordBlocks[i]) + str(self.recordRooms[i])}\t\t"
                                  f"{str(self.recordTelNumber[i])}\t\t   {str(self.recordDates[i])}\t\t{str(self.recordTimes[i])}\n")

        self.textPlace.config(state='disabled')

    #Открытие файла
    def openExcelFile(self):
        filePath = tk.filedialog.askopenfilename(title="Выберите файл",
                                                 filetypes=[("Excel файлы", "*.xls"), ("Excel файлы", "*.xlsx")])
        # Проверка, выбрал ли пользователь файл или нажал "Отмена"
        if not filePath:
            return

        try:
            self.wb = load_workbook(filePath)
            self.activeList = self.wb.active
        except Exception as e:
            tk.messagebox.showerror("Ошибка файла", f"Не удалось открыть файл: {str(e)}")
            return

        self.readFromExcel()

    def infoToDocx(self):
        self.filePath = 'wordFiles//Word1.docx'
        self.document = Document()

        # Добавляем заголовок документа
        self.document.add_heading("Данные о записях на дежурство", 0)

        # Добавляем таблицу в документ
        table = self.document.add_table(rows=1, cols=8)

        # Определяем заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Фамилия'
        hdr_cells[1].text = 'Имя'
        hdr_cells[2].text = 'Блок'
        hdr_cells[3].text = 'Комната'
        hdr_cells[4].text = 'Номер'
        hdr_cells[5].text = 'Дата'
        hdr_cells[6].text = 'Год'
        hdr_cells[7].text = 'Время'

        # Заполняем таблицу данными
        for i in range(len(self.recordBlocks)):
            row_cells = table.add_row().cells
            rowData = [
                self.recordSecondNames[i],
                self.recordFirstNames[i],
                self.recordBlocks[i],
                self.recordRooms[i],
                self.recordTelNumber[i],
                self.recordDates[i],
                self.recordYears[i],  # Добавляем год
                self.recordTimes[i]
            ]

            if all(isinstance(item, (int, float, str)) for item in rowData):
                for j, cell in enumerate(row_cells):
                    cell.text = str(rowData[j])

        # Сохраняем документ
        self.document.save(self.filePath)

        tk.messagebox.showinfo("Сохранение", f"Данные успешно сохранены!\n{self.filePath}")

    #Функция отправки файла в бот
    def sendFileToBot(self):

        self.infoToDocx()

        filePath = 'wordFiles//Word1.docx'
        self.telegramBot.sendDocument(self.chatID, filePath)

    def checkIsEntry(self):

        if (self.comboboxSort.get()):

            self.comboboxSort.config(state='readonly')

            return True

        else:
            return False


    # Функция которая проверяет не одна ли строка в тексте (для сортировки нужно минимум 2)
    def checkIsNoOneLine(self):

        numberLines = int(self.textPlace.index("end-1c").split(".")[0])

        if (numberLines >= 5):

            return True

        else:

            return False

    # Функция сортировки
    def funcOfSort(self):

        self.selectedSort = self.comboboxSort.get()

        resChecking = self.checkIsEntry()

        if not resChecking:
            # parent для того, чтобы при вызове ошибки sortWindow не закрывалось
            messagebox.showerror("Ошибка", "Выберите способ сортировки!")

        # Здесь сама сортировка через else (просто напиши else и там вызывай функции сортировки)
        elif resChecking:

            NoOneLine = self.checkIsNoOneLine()

            # Здесь вызываем функции сортировки
            if NoOneLine:

                if self.selectedSort == "По возрастанию блоков":

                    self.sortUpBlocks()

                elif self.selectedSort == "По убыванию блоков":

                    self.sortDownBlocks()

                elif self.selectedSort == "По фамилиям (от А до Я)":

                    self.sortByAlphabetUp()

                elif self.selectedSort == "По фамилиям (от Я до А)":
                    self.sortByAlphabetDown()

                elif self.selectedSort == "По дате":

                    self.sortByData()

            if not NoOneLine:

                tk.messagebox.showerror("Ошибка", "Недостаточно записей (минимум 2)")

    # Функция сортировки по возрастанию блоков
    def sortUpBlocks(self):
        self.removeArrowAndHighlight(self.selected_index)
        # Создание списка кортежей (блок, индекс)
        combined_list = list(enumerate(self.recordBlocks))
        # Сортировка по блокам
        sorted_combined_list = sorted(combined_list, key=lambda x: x[1])

        # Очищаем textPlace
        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', "\n")

        # Вставляем отсортированные строки в textPlace
        for index, _ in sorted_combined_list:
            fullString = self.recordFullStrings[index].strip()  # Удаление лишних пробелов и переносов строк
            self.textPlace.insert(tk.END, fullString + "\n")

        self.textPlace.config(state='disabled')

    def sortDownBlocks(self):
        self.removeArrowAndHighlight(self.selected_index)
        # Создание списка кортежей (блок, индекс)
        combined_list = list(enumerate(self.recordBlocks))
        # Сортировка по блокам
        sorted_combined_list = sorted(combined_list, key=lambda x: x[1], reverse=True)

        # Очищаем textPlace
        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', "\n")

        # Вставляем отсортированные строки в textPlace
        for index, _ in sorted_combined_list:
            fullString = self.recordFullStrings[index].strip()  # Удаление лишних пробелов и переносов строк
            self.textPlace.insert(tk.END, fullString + "\n")

        self.textPlace.config(state='disabled')

    def sortByAlphabetUp(self):
        self.removeArrowAndHighlight(self.selected_index)
        # Сортируем полный список строк по алфавиту
        sortedAlphabetical = sorted(self.recordFullStrings)

        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', "\n")

        for string in sortedAlphabetical:
            self.textPlace.insert(tk.END, string.strip() + "\n")  # Удаление лишних пробелов и переносов строк

        self.textPlace.config(state='disabled')

    def sortByAlphabetDown(self):
        self.removeArrowAndHighlight(self.selected_index)
        # Сортируем полный список строк по алфавиту в обратном порядке
        sortedAlphabetical = sorted(self.recordFullStrings, reverse=True)

        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', "\n")

        for string in sortedAlphabetical:
            self.textPlace.insert(tk.END, string.strip() + "\n")  # Удаление лишних пробелов и переносов строк

        self.textPlace.config(state='disabled')

    def sortByData(self):
        self.removeArrowAndHighlight(self.selected_index)
        # Создание списка кортежей (месяц, день, индекс)
        combined_list = list(enumerate(zip(self.recordMonthes, self.recordDays)))

        # Преобразование данных к типу int для корректной сортировки
        combined_list = [(i, (int(month), int(day))) for i, (month, day) in combined_list]

        # Сортировка по месяцу и дню
        sorted_combined_list = sorted(combined_list, key=lambda x: (x[1][0], x[1][1]))

        # Очищаем textPlace
        self.textPlace.config(state='normal')
        self.textPlace.delete('3.0', tk.END)
        self.textPlace.insert('2.0', "\n")

        # Вставляем отсортированные строки в textPlace
        for index, _ in sorted_combined_list:
            if index < len(self.recordFullStrings):
                fullString = self.recordFullStrings[index].strip()  # Удаление лишних пробелов и переносов строк
                self.textPlace.insert(tk.END, fullString + "\n")
            else:
                print(
                    f"Индекс {index} выходит за пределы списка recordFullStrings длиной {len(self.recordFullStrings)}")  # Отладочное сообщение

        self.textPlace.config(state='disabled')

    # ---------------------------------------------------------------------------------------------------------------ФУНКЦИИ

#Класс окна помощи
class HelpWindow():

    def __init__(self, frame, parent):

        self.parent = parent

        self.helpWindow = tk.Toplevel(frame)

        self.helpWindow.title("Использование")
        self.helpWindow.geometry("600x400")
        self.helpWindow.geometry("+630+270")
        self.helpWindow.resizable(width = False, height=False)

        self.frameHelp1 = tk.Frame(self.helpWindow, width = 600, height = 600)
        self.frameHelp1.place(x = 0, y = 0)

        self.frameHelp2 = tk.Frame(self.helpWindow, width=600, height=600)
        self.frameHelp3 = tk.Frame(self.helpWindow, width=600, height=600)
        self.frameHelp4 = tk.Frame(self.helpWindow, width=600, height=600)
        self.frameHelp5 = tk.Frame(self.helpWindow, width=600, height=600)
        self.frameHelp6 = tk.Frame(self.helpWindow, width=600, height=600)
        self.frameHelp7 = tk.Frame(self.helpWindow, width=600, height=600)

        self.canvasHelp1 = tk.Canvas(self.frameHelp1, width=130, height=600, bg='#3469ED')
        self.canvasHelp1.place(x = 470, y = 0)

        self.canvasHelp2 = tk.Canvas(self.frameHelp2, width=130, height=600, bg='#3469ED')
        self.canvasHelp2.place(x=470, y=0)

        self.canvasHelp3 = tk.Canvas(self.frameHelp3, width=130, height=600, bg='#3469ED')
        self.canvasHelp3.place(x=470, y=0)

        self.canvasHelp4 = tk.Canvas(self.frameHelp4, width=130, height=600, bg='#3469ED')
        self.canvasHelp4.place(x=470, y=0)

        self.canvasHelp5 = tk.Canvas(self.frameHelp5, width=130, height=600, bg='#3469ED')
        self.canvasHelp5.place(x=470, y=0)

        self.canvasHelp6 = tk.Canvas(self.frameHelp6, width=130, height=600, bg='#3469ED')
        self.canvasHelp6.place(x=470, y=0)

        self.imageWrite = tk.PhotoImage(file = 'images//icons//Write.png')
        self.labelForWrite = MyLabel(self.frameHelp1, image = self.imageWrite, bg = "#3469ED", x = 490, y = 135)
        self.labelForWrite.image = self.imageWrite

        self.imageClear = tk.PhotoImage(file='images//icons//Clear.png')
        self.labelForClear = MyLabel(self.frameHelp2, image=self.imageClear, bg="#3469ED", x=490, y=135)
        self.labelForClear.image = self.imageClear

        self.imageDelete = tk.PhotoImage(file='images//icons//Delete.png')
        self.labelForDelete = MyLabel(self.frameHelp3, image=self.imageDelete, bg="#3469ED", x=490, y=135)
        self.labelForDelete.image = self.imageDelete

        self.imageExcel = tk.PhotoImage(file='images//icons//Excel.png')
        self.labelForExcel = MyLabel(self.frameHelp4, image=self.imageExcel, bg="#3469ED", x=490, y=135)
        self.labelForExcel.image = self.imageExcel

        self.imageWord = tk.PhotoImage(file='images//icons//Word.png')
        self.labelForWord = MyLabel(self.frameHelp5, image=self.imageWord, bg="#3469ED", x=490, y=135)
        self.labelForWord.image = self.imageWord

        self.imageSend = tk.PhotoImage(file='images//icons//Send.png')
        self.labelForSend = MyLabel(self.frameHelp6, image=self.imageSend, bg="#3469ED", x=490, y=135)
        self.labelForSend.image = self.imageSend


        self.labelHelp1 = tk.Label(self.frameHelp1, text = "Кнопка 'Записаться'", font = ("Montserrat", 13, 'bold'))
        self.labelHelp1.place(x = 170, y = 25)

        self.labelHelp11 = tk.Label(self.frameHelp1, text="Данная кнопка позволяет пользователю записаться на дежурство\n"
                                                         "в случае заполнения всех полей корректными данными.\n\n"
                                                         "При незаполнении каких-либо полей или некорректном их заполнении\n"
                                                         "Программа выдаст пользователю соответствующую ошибку.\n\n"
                                                         "Результат записи можно увидеть в главном окне. Все записи\n"
                                                         "будут расположены упорядоченно и удобно для чтения.\n\n"
                                                          "Доступ : студент, комендант", font=("Montserrat", 10), justify = "left")
        self.labelHelp11.place(x=30, y=100)

        self.labelHelp2 = tk.Label(self.frameHelp2, text="Кнопка 'Очистить ввод'", font=("Montserrat", 13, 'bold'))
        self.labelHelp2.place(x=170, y=25)

        self.labelHelp22 = tk.Label(self.frameHelp2, text = "Данная кнопка позволяет очищать все поля записей\n"
                                                            "независимо от корректности ввода данных пользователем.\n\n"
                                                            "Доступ : студент, комендант", font=("Montserrat", 10), justify = "left")
        self.labelHelp22.place(x = 30, y = 100)

        self.labelHelp3 = tk.Label(self.frameHelp3, text="Кнопка 'Удалить данные'", font=("Montserrat", 13, 'bold'))
        self.labelHelp3.place(x=170, y=25)

        self.labelHelp33 = tk.Label(self.frameHelp3, text="Данная кнопка позволяет очищать все записанные данные в\n"
                                                          "приложении без возможности возврата.\n\nДоступ : комендант",
                                    font=("Montserrat", 10), justify="left")
        self.labelHelp33.place(x=30, y=100)

        self.labelHelp4 = tk.Label(self.frameHelp4, text="Кнопка 'Сохранить в Excel'", font=("Montserrat", 13, 'bold'))
        self.labelHelp4.place(x=170, y=25)

        self.labelHelp44 = tk.Label(self.frameHelp4, text="Данная кнопка позволяет записывать данные в Excel файл.\n"
                                                          "Пользователю предлагается выбрать самим файл для записи.\n"
                                                          "Если файла по умолчанию нет, он создастся в папке с проектом.\n\n"
                                                          "Доступ : студент, комендант", font=("Montserrat", 10), justify="left")
        self.labelHelp44.place(x=30, y=100)

        self.labelHelp5 = tk.Label(self.frameHelp5, text="Кнопка 'Сохранить в Word'", font=("Montserrat", 13, 'bold'))
        self.labelHelp5.place(x=170, y=25)

        self.labelHelp55 = tk.Label(self.frameHelp5, text="Данная кнопка позволяет записывать данные в Word файл.\n"
                                                          "Если файла по умолчанию нет, он создастся в папке с проектом.\n\n"
                                                          "Доступ : студент, комендант", font=("Montserrat", 10), justify="left")
        self.labelHelp55.place(x=30, y=100)

        self.labelHelp6 = tk.Label(self.frameHelp6, text="Кнопка 'Отправить коменданту'", font=("Montserrat", 13, 'bold'))
        self.labelHelp6.place(x=170, y=25)

        self.labelHelp66 = tk.Label(self.frameHelp6, text="Данная кнопка позволяет отправлять файл формата .docx\n"
                                                          "с текущими записанными данными о дежурстве в телеграмм - бот\n\n"
                                                          "Доступ : студент", font=("Montserrat", 10), justify="left")
        self.labelHelp66.place(x=30, y=100)

        self.addNavigationButtons()

    def addNavigationButtons(self):

        MyButton(self.frameHelp1, text="Далее", width=10, height=2, x=350, y=300, command=self.fromHelp1ToHelp2)
        MyButton(self.frameHelp2, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp2ToHelp1)

        MyButton(self.frameHelp2, text="Далее", width=10, height=2, x=350, y=300, command=self.fromHelp2ToHelp3)
        MyButton(self.frameHelp2, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp2ToHelp1)

        MyButton(self.frameHelp3, text="Далее", width=10, height=2, x=350, y=300, command=self.fromHelp3ToHelp4)
        MyButton(self.frameHelp3, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp3ToHelp2)

        MyButton(self.frameHelp4, text="Далее", width=10, height=2, x=350, y=300, command=self.fromHelp4ToHelp5)
        MyButton(self.frameHelp4, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp4ToHelp3)

        MyButton(self.frameHelp5, text="Далее", width=10, height=2, x=350, y=300, command=self.fromHelp5ToHelp6)
        MyButton(self.frameHelp5, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp5ToHelp4)

        MyButton(self.frameHelp6, text="Выход", width=10, height=2, x=350, y=300, command=self.closeWindowHelp)
        MyButton(self.frameHelp6, text="Назад", width=10, height=2, x=50, y=300, command=self.fromHelp6ToHelp5)


    def fromHelp1ToHelp2(self):
        self.frameHelp1.pack_forget()
        self.frameHelp2.pack()

    def fromHelp2ToHelp3(self):
        self.frameHelp2.pack_forget()
        self.frameHelp3.pack()

    def fromHelp2ToHelp1(self):
        self.frameHelp2.pack_forget()
        self.frameHelp1.pack()

    def fromHelp3ToHelp2(self):
        self.frameHelp3.pack_forget()
        self.frameHelp2.pack()

    def fromHelp3ToHelp4(self):
        self.frameHelp3.pack_forget()
        self.frameHelp4.pack()

    def fromHelp4ToHelp5(self):
        self.frameHelp4.pack_forget()
        self.frameHelp5.pack()

    def fromHelp4ToHelp3(self):
        self.frameHelp4.pack_forget()
        self.frameHelp3.pack()

    def fromHelp5ToHelp6(self):
        self.frameHelp5.pack_forget()
        self.frameHelp6.pack()

    def fromHelp5ToHelp4(self):
        self.frameHelp5.pack_forget()
        self.frameHelp4.pack()

    def fromHelp6ToHelp5(self):
        self.frameHelp6.pack_forget()
        self.frameHelp5.pack()

    def closeWindowHelp(self):
        self.helpWindow.destroy()

#Класс для работы с Excel
class WorkExcel():

    def __init__(self, parent, filepath):

        self.parent = parent

        self.filepath = filepath

        self.wb = Workbook()

        self.activeList = self.wb.active

        self.activeList.title = "Sheet1"

        self.boldFont = Font(bold=True)

        self.activeList['A1'] = "Фамилия"
        self.activeList['A1'].font = self.boldFont

        self.activeList['B1'] = "Имя"
        self.activeList['B1'].font = self.boldFont

        self.activeList['C1'] = "Блок"
        self.activeList['C1'].font = self.boldFont

        self.activeList['D1'] = "Комната"
        self.activeList['D1'].font = self.boldFont

        self.activeList['E1'] = "Номер"
        self.activeList['E1'].font = self.boldFont

        self.activeList['F1'] = "Дата"
        self.activeList['F1'].font = self.boldFont

        self.activeList['G1'] = "Время"
        self.activeList['G1'].font = self.boldFont

        self.wb.save(self.filepath)

    #Функция записывающая данные в таблицу
    def infoInFile(self, filepath):

        self.filepath = filepath

        for i in range(len(self.parent.recordBlocks)):

            rowData = [self.parent.recordSecondNames[i], self.parent.recordFirstNames[i],self.parent.recordBlocks[i],
                                    self.parent.recordRooms[i], self.parent.recordTelNumber[i], self.parent.recordDates[i], self.parent.recordTimes[i]]

            if all(isinstance(item, (int, float, str)) for item in rowData):

                self.activeList.append(rowData)

                self.wb.save(self.filepath)

# Класс работы с телеграм-ботом
class MyTelegramBot:

    def __init__(self, parent):

        self.parent = parent

        self.bot = telebot.TeleBot('7931036017:AAF-C7LUTnueZ1Mgftg8uw1j0YjpH76rzZ0')

        @self.bot.message_handler(commands=['start'])
        def startFunc(message):
            if self.parent.checkInternetConnection():
                self.bot.send_message(message.chat.id, f"Здравствуйте {message.from_user.first_name}. Данный бот разработан студентом БНТУ ФИТР гр. 10701123 Жоровым Е.А для отправки файлов с данными о записях на дежурство в общежитии")

        @self.bot.message_handler(commands=['about'])
        def abtProgram(message):
            if self.parent.checkInternetConnection():
                self.bot.send_message(message.chat.id, "Здравствуйте. Данный бот разработан студентом БНТУ ФИТР гр. 10701123 Жоровым Е.А для отправки файлов с данными о записях на дежурство в общежитии")


        @self.bot.message_handler(commands=['id'])
        def getID(message):
            if self.parent.checkInternetConnection():
                self.bot.send_message(message.chat.id, message.chat.id)

    def sendDocument(self, chat_id, file_path):
        if self.parent.checkInternetConnection():
            with open(file_path, 'rb') as file:
                self.bot.send_document(chat_id, file)
        else:
            tk.messagebox.showerror("Ошибка", "Отсутсвует подключение к интернету")

    def runBot(self):
        self.bot.infinity_polling()

def runTelegramBot(bot_instance):
    bot_instance.runBot()

app = MyApp()

# Создаем и запускаем поток для бота
bot_thread = threading.Thread(target=runTelegramBot, args=(app.telegramBot,))

bot_thread.daemon = True # Задает поток как фоновый (закрывается при завершении основного потока)
bot_thread.start()

app.mainloop()

